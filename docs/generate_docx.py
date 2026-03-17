"""Generate Word documents for UK and DE/NL Market Data Subscriptions."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shading_elem)


def style_header_row(table, color_hex="1F3864"):
    """Style the first row of a table as a header."""
    for cell in table.rows[0].cells:
        set_cell_shading(cell, color_hex)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    style_header_row(table)

    # Alternating row shading
    for r_idx in range(1, len(table.rows)):
        if r_idx % 2 == 0:
            for cell in table.rows[r_idx].cells:
                set_cell_shading(cell, "E8EDF3")

    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    return h


def add_bullet(doc, label, text):
    """Add a bold label + normal text as a paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    run_b = p.add_run(f"{label}: ")
    run_b.bold = True
    run_b.font.size = Pt(10)
    run_n = p.add_run(text)
    run_n.font.size = Pt(10)


def add_data_feed_section(doc, number, title, items):
    """Add a numbered data feed with bullet-point details."""
    add_heading(doc, f"{number} {title}", level=3)
    for label, text in items:
        add_bullet(doc, label, text)


# ============================================================
# DOCUMENT 1: UK Market Rate Data Subscriptions
# ============================================================

def generate_uk_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("UK Market Rate Data Subscriptions", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # Metadata
    meta = doc.add_paragraph()
    for label, val in [
        ("Asset", "Northwold Solar Farm (Hall Farm) \u2014 8.4 MWh / 7.5 MW BESS"),
        ("Aggregator", "GridBeyond"),
        ("Prepared", "February 2026"),
    ]:
        run_b = meta.add_run(f"{label}: ")
        run_b.bold = True
        run_b.font.size = Pt(10)
        run_n = meta.add_run(f"{val}\n")
        run_n.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ---- Section 1: Wholesale ----
    add_heading(doc, "1. Wholesale Electricity Market Prices", level=1)

    add_data_feed_section(doc, "1.1", "EPEX Day-Ahead Auction Price", [
        ("What", "Next-day electricity price set via daily auction on EPEX SPOT exchange"),
        ("Why needed", "Primary benchmark for arbitrage strategy \u2014 determines when to charge (low price) and discharge (high price)"),
        ("Frequency", "Half-hourly (48 values per day)"),
        ("Unit", "\u00a3/MWh"),
        ("Current delivery", "Included in GridBeyond monthly report"),
        ("Direct source", "EPEX SPOT Market Data Services"),
        ("Direct cost", "Quote required (contact marketdata@epexspot.com)"),
        ("Free alternative", "Elexon BMRS Market Index Data (DA HH reference price)"),
    ])

    add_data_feed_section(doc, "1.2", "EPEX 30-Minute Day-Ahead Price", [
        ("What", "EPEX SPOT 30-minute ahead product price \u2014 a shorter-horizon variant of the DA auction"),
        ("Why needed", "Captures intra-day price movements closer to delivery; used for short-term dispatch decisions"),
        ("Frequency", "Half-hourly"),
        ("Unit", "\u00a3/MWh"),
        ("Current delivery", "Included in GridBeyond monthly report"),
        ("Direct source", "EPEX SPOT"),
        ("Direct cost", "Quote required"),
    ])

    add_data_feed_section(doc, "1.3", "GB-ISEM Intraday Auction 1 (IDA1) Price", [
        ("What", "Same-day auction price on the GB-ISEM (Integrated Single Electricity Market) intraday market"),
        ("Why needed", "Second arbitrage opportunity \u2014 prices may differ from day-ahead, offering improved trading margins"),
        ("Frequency", "Half-hourly"),
        ("Unit", "\u00a3/MWh"),
        ("Current delivery", "Included in GridBeyond monthly report"),
        ("Direct source", "EPEX SPOT / SEMOpx"),
        ("Direct cost", "Bundled with EPEX subscription"),
    ])

    add_data_feed_section(doc, "1.4", "Intraday Continuous (IDC) Price", [
        ("What", "Continuous electronic trading price on EPEX intraday market \u2014 trades executed up to gate closure"),
        ("Why needed", "Captures real-time price spikes and short-notice trading opportunities"),
        ("Frequency", "Half-hourly (aggregated from continuous trades)"),
        ("Unit", "\u00a3/MWh"),
        ("Current delivery", "Included in GridBeyond monthly report"),
        ("Direct source", "EPEX SPOT / Nord Pool"),
        ("Direct cost", "Nord Pool Intraday: ~\u20ac5,000/year; EPEX: quote required"),
    ])

    add_data_feed_section(doc, "1.5", "Day-Ahead Half-Hourly Reference Price (DA HH)", [
        ("What", "Elexon\u2019s official day-ahead reference price \u2014 market index used in settlement"),
        ("Why needed", "Independent benchmark for validating EPEX prices and used in optimisation engine"),
        ("Frequency", "Half-hourly"),
        ("Unit", "\u00a3/MWh"),
        ("Current delivery", "Included in GridBeyond monthly report"),
        ("Direct source", "Elexon BMRS API"),
        ("Direct cost", "FREE (no authentication required)"),
    ])

    add_data_feed_section(doc, "1.6", "System Sell Price (SSP)", [
        ("What", "Price the grid pays generators/batteries when the system is long (excess supply)"),
        ("Why needed", "Determines revenue for unplanned exports; key input for imbalance analysis"),
        ("Frequency", "Half-hourly"),
        ("Unit", "\u00a3/MWh"),
        ("Current delivery", "Included in GridBeyond monthly report"),
        ("Direct source", "Elexon BMRS API"),
        ("Direct cost", "FREE"),
    ])

    add_data_feed_section(doc, "1.7", "System Buy Price (SBP)", [
        ("What", "Price batteries/generators pay when the system is short (deficit supply)"),
        ("Why needed", "Determines cost of unplanned imports; critical for imbalance risk management"),
        ("Frequency", "Half-hourly"),
        ("Unit", "\u00a3/MWh"),
        ("Current delivery", "Included in GridBeyond monthly report"),
        ("Direct source", "Elexon BMRS API"),
        ("Direct cost", "FREE"),
    ])

    # ---- Section 2: Ancillary ----
    add_heading(doc, "2. Ancillary Services Clearing Prices", level=1)

    p = doc.add_paragraph(
        "All ancillary services below are procured by NESO (National Energy System Operator) via competitive tenders. "
        "Clearing prices determine the revenue rate per MW of availability provided."
    )
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True

    ancillary_services = [
        ("2.1", "Static Firm Frequency Response (SFFR) Clearing Price", [
            ("What", "Auction clearing price for providing continuous frequency support to the grid"),
            ("Why needed", "SFFR is currently the largest single revenue stream for the asset \u2014 essential for strategy comparison (hold SFFR vs. trade wholesale)"),
            ("Frequency", "Per EFA block (6 blocks \u00d7 4 hours = 24 hours)"),
            ("Unit", "\u00a3/MW/h"),
            ("Current delivery", "Included in GridBeyond monthly report"),
            ("Direct source", "NESO Data Portal"),
            ("Direct cost", "FREE"),
        ]),
        ("2.2", "Dynamic Containment Low (DCL) Clearing Price", [
            ("What", "Price for providing low-frequency containment response (grid frequency drops below 49.5 Hz)"),
            ("Why needed", "DCL is a key revenue alternative to wholesale trading; pricing trends indicate market saturation"),
            ("Frequency", "Per EFA block"),
            ("Unit", "\u00a3/MW/h"),
            ("Direct source", "NESO Data Portal"),
            ("Direct cost", "FREE"),
        ]),
        ("2.3", "Dynamic Containment High (DCH) Clearing Price", [
            ("What", "Price for providing high-frequency containment response (grid frequency rises above 50.5 Hz)"),
            ("Why needed", "Paired with DCL \u2014 both needed to evaluate full DC revenue potential"),
            ("Frequency", "Per EFA block"),
            ("Unit", "\u00a3/MW/h"),
            ("Direct source", "NESO Data Portal"),
            ("Direct cost", "FREE"),
        ]),
        ("2.4", "Dynamic Modulation Low (DML) Clearing Price", [
            ("What", "Price for enhanced frequency response at moderate deviations (below 49.75 Hz)"),
            ("Why needed", "Newer product \u2014 tracks evolving ancillary services market and revenue diversification options"),
            ("Frequency", "Per EFA block"),
            ("Unit", "\u00a3/MW/h"),
            ("Direct source", "NESO Data Portal"),
            ("Direct cost", "FREE"),
        ]),
        ("2.5", "Dynamic Modulation High (DMH) Clearing Price", [
            ("What", "Price for enhanced frequency response at moderate deviations (above 50.25 Hz)"),
            ("Why needed", "Paired with DML for full DM evaluation"),
            ("Frequency", "Per EFA block"),
            ("Unit", "\u00a3/MW/h"),
            ("Direct source", "NESO Data Portal"),
            ("Direct cost", "FREE"),
        ]),
        ("2.6", "Dynamic Regulation Low (DRL) Clearing Price", [
            ("What", "Price for providing continuous active power regulation for low-frequency events"),
            ("Why needed", "Completes the frequency response product suite \u2014 needed for comprehensive ancillary revenue analysis"),
            ("Frequency", "Per EFA block"),
            ("Unit", "\u00a3/MW/h"),
            ("Direct source", "NESO Data Portal"),
            ("Direct cost", "FREE"),
        ]),
        ("2.7", "Dynamic Regulation High (DRH) Clearing Price", [
            ("What", "Price for providing continuous active power regulation for high-frequency events"),
            ("Why needed", "Paired with DRL for full DR evaluation"),
            ("Frequency", "Per EFA block"),
            ("Unit", "\u00a3/MW/h"),
            ("Direct source", "NESO Data Portal"),
            ("Direct cost", "FREE"),
        ]),
    ]

    for num, title, items in ancillary_services:
        add_data_feed_section(doc, num, title, items)

    # ---- Section 3: Imbalance ----
    add_heading(doc, "3. Imbalance & Settlement Data", level=1)

    add_data_feed_section(doc, "3.1", "Imbalance Settlement Prices", [
        ("What", "Half-hourly settlement prices applied to any energy volume that deviates from contracted position"),
        ("Why needed", "Imbalance charges are a significant cost line \u2014 understanding exposure is critical for risk management"),
        ("Frequency", "Half-hourly"),
        ("Unit", "\u00a3/MWh"),
        ("Current delivery", "Revenue and charges included in GridBeyond monthly report"),
        ("Direct source", "Elexon BMRS API (System Price data)"),
        ("Direct cost", "FREE"),
    ])

    # ---- Section 4: Benchmarking ----
    add_heading(doc, "4. Benchmarking & Analytics Subscriptions", level=1)

    add_data_feed_section(doc, "4.1", "Modo Energy \u2014 BESS Revenue Benchmarks", [
        ("What", "Third-party analytics platform providing GB BESS market indices, revenue benchmarks, and asset-level performance data"),
        ("Why needed", "Enables comparison of Northwold performance against industry peers; provides TB spread indices (TB1, TB2, TB3) and revenue forecasts"),
        ("Frequency", "Daily / Monthly reports"),
        ("Current status", "Referenced in dashboard benchmarks; not a live integration"),
        ("Cost tiers", "Free: limited overview | Plus: basic analytics | Pro: ~\u00a38,000\u2013\u00a315,000/year | Business: ~\u00a315,000\u2013\u00a330,000/year | Enterprise: ~\u00a330,000\u2013\u00a360,000/year"),
    ])

    add_data_feed_section(doc, "4.2", "Nord Pool \u2014 UK Day-Ahead & Intraday Market Data", [
        ("What", "Alternative market data provider for GB day-ahead (N2EX) and intraday trading data"),
        ("Why needed", "Potential direct feed if moving away from reliance on GridBeyond report; provides real-time and historical price data"),
        ("Frequency", "Real-time / daily"),
        ("Cost", "~\u20ac1,200/year (UK Day-Ahead) + ~\u20ac5,000/year (Intraday Trades) = ~\u20ac6,200/year (~\u00a35,300/year)"),
        ("Current status", "Not currently subscribed \u2014 data comes via GridBeyond"),
    ])

    # ---- Section 5: Cost Summary ----
    add_heading(doc, "5. Cost Summary", level=1)

    add_table(doc,
        ["Data Source", "Annual Cost", "Status"],
        [
            ["GridBeyond Monthly Report (all wholesale + ancillary data)", "\u00a30 (included in 5% revenue share)", "Active"],
            ["Elexon BMRS API (SSP, SBP, DA HH, settlement)", "\u00a30 (free public API)", "Available"],
            ["NESO Data Portal (all ancillary clearing prices)", "\u00a30 (free public data)", "Available"],
            ["Nord Pool (UK DA + Intraday direct feed)", "~\u00a35,300/year", "Optional"],
            ["EPEX SPOT (direct market data feed)", "Quote required", "Optional"],
            ["Modo Energy (BESS benchmarking \u2014 Pro tier)", "~\u00a38,000\u2013\u00a315,000/year", "Optional"],
            ["Total (current operating cost)", "\u00a30/year", ""],
            ["Total (with benchmarking + direct feeds)", "\u00a313,300\u2013\u00a320,300+/year", ""],
        ],
    )

    # ---- Section 6: Notes ----
    add_heading(doc, "6. Notes", level=1)

    notes = [
        "GridBeyond is the primary data delivery mechanism \u2014 all 7 wholesale prices and 7 ancillary clearing prices arrive bundled in their monthly Excel report at no additional cost (covered by the 5% aggregator revenue share).",
        "Direct subscriptions become relevant if: moving to real-time or near-real-time dashboard updates; reducing dependency on GridBeyond for data accuracy/timeliness; implementing live trading decision support; or if the GridBeyond contract changes.",
        "Elexon BMRS and NESO are always free \u2014 these are public UK energy data platforms with open APIs, suitable for independent verification of GridBeyond data.",
        "Data lag consideration: GridBeyond reports arrive ~2 weeks after month-end. Direct API subscriptions would reduce this to near-real-time.",
    ]
    for i, note in enumerate(notes, 1):
        p = doc.add_paragraph(style="List Number")
        p.text = note
        for run in p.runs:
            run.font.size = Pt(10)

    doc.save(r"C:\repos\bess-dashboard\docs\UK_Market_Rate_Data_Subscriptions.docx")
    print("UK document saved.")


# ============================================================
# DOCUMENT 2: DE/NL Market Rate Data Subscriptions
# ============================================================

def generate_de_nl_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("Market Rate Data Subscriptions \u2014 Germany & Netherlands", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # Metadata
    meta = doc.add_paragraph()
    for label, val in [
        ("Purpose", "Data requirements for extending the BESS dashboard to German and Dutch assets"),
        ("Reference", "See UK Market Rate Data Subscriptions document for existing UK requirements"),
        ("Prepared", "February 2026"),
    ]:
        run_b = meta.add_run(f"{label}: ")
        run_b.bold = True
        run_b.font.size = Pt(10)
        run_n = meta.add_run(f"{val}\n")
        run_n.font.size = Pt(10)

    doc.add_paragraph()

    # ---- Structural Differences ----
    add_heading(doc, "Key Structural Differences from UK", level=1)

    add_table(doc,
        ["Feature", "UK (GB)", "Germany (DE)", "Netherlands (NL)"],
        [
            ["Settlement period", "30 min (48/day)", "15 min (96/day)", "15 min (96/day)"],
            ["Currency", "GBP", "EUR", "EUR"],
            ["TSO(s)", "NESO (1)", "50Hertz, Amprion, TenneT, TransnetBW (4)", "TenneT NL (1)"],
            ["Bidding zone", "GB", "DE-LU", "NL"],
            ["Exchange", "EPEX SPOT / Nord Pool", "EPEX SPOT", "EPEX SPOT"],
            ["Imbalance pricing", "Dual (SSP / SBP)", "Single (reBAP)", "Hybrid single/dual"],
            ["Frequency response", "SFFR, DC, DM, DR (GB-specific)", "FCR, aFRR, mFRR (EU std)", "FCR, aFRR, mFRR (EU std)"],
            ["Intraday auctions/day", "1 (IDA1)", "3 (IDA1, IDA2, IDA3)", "3 (IDA1, IDA2, IDA3)"],
            ["Intraday gate closure", "30 min before delivery", "5 min before delivery", "5 min before delivery"],
        ],
    )

    p = doc.add_paragraph(
        "\nDashboard implication: The move from 48 to 96 periods per day doubles the data volume. "
        "The ancillary services module needs a complete redesign \u2014 DE/NL use FCR/aFRR/mFRR instead of UK\u2019s SFFR/DC/DM/DR products."
    )
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True

    # ==============================
    # PART 1: GERMANY
    # ==============================
    add_heading(doc, "PART 1: GERMANY (DE)", level=1)

    # -- DE Wholesale --
    add_heading(doc, "1. Wholesale Electricity Prices \u2014 Germany", level=2)

    de_wholesale = [
        ("1.1", "EPEX Day-Ahead Auction Price (DE-LU)", [
            ("What", "Next-day electricity price from the EPEX SPOT day-ahead auction for the Germany-Luxembourg bidding zone"),
            ("Why needed", "Primary benchmark for charge/discharge arbitrage decisions"),
            ("Frequency", "Quarter-hourly (96 values per day, since October 2025)"),
            ("Unit", "\u20ac/MWh"),
            ("Free source", "SMARD (Bundesnetzagentur) or ENTSO-E Transparency Platform"),
            ("Paid source", "EPEX SPOT (quote required)"),
            ("Cost", "FREE via SMARD / ENTSO-E"),
        ]),
        ("1.2", "EPEX Intraday Auction 1 (IDA1) Price \u2014 DE", [
            ("What", "First same-day auction result \u2014 provides updated pricing after day-ahead"),
            ("Why needed", "Second trading opportunity to refine positions before delivery"),
            ("Frequency", "Quarter-hourly"),
            ("Unit", "\u20ac/MWh"),
            ("Gate closure", "15:00 CET on D-1"),
            ("Source", "EPEX SPOT (paid subscription); partial data on SMARD"),
            ("Cost", "Bundled with EPEX SPOT subscription (quote required)"),
        ]),
        ("1.3", "EPEX Intraday Auction 2 (IDA2) Price \u2014 DE", [
            ("What", "Second same-day auction \u2014 evening auction for full next day"),
            ("Why needed", "Further position refinement; captures late-day forecast changes"),
            ("Frequency", "Quarter-hourly"),
            ("Unit", "\u20ac/MWh"),
            ("Gate closure", "22:00 CET on D-1"),
            ("Source", "EPEX SPOT (paid subscription)"),
            ("Cost", "Bundled with EPEX SPOT subscription"),
        ]),
        ("1.4", "EPEX Intraday Auction 3 (IDA3) Price \u2014 DE", [
            ("What", "Third intraday auction \u2014 morning of delivery day for afternoon/evening periods"),
            ("Why needed", "Captures same-day demand and generation updates"),
            ("Frequency", "Quarter-hourly (covers 12:00\u201324:00 delivery)"),
            ("Unit", "\u20ac/MWh"),
            ("Gate closure", "10:00 CET on delivery day"),
            ("Source", "EPEX SPOT (paid subscription)"),
            ("Cost", "Bundled with EPEX SPOT subscription"),
        ]),
        ("1.5", "Intraday Continuous (IDC) Price \u2014 DE", [
            ("What", "Continuous electronic trading price on EPEX intraday market \u2014 trades executed up to 5 minutes before delivery"),
            ("Why needed", "Captures real-time price volatility and short-notice opportunities; 5-minute gate closure is much shorter than UK\u2019s 30-minute"),
            ("Frequency", "Quarter-hourly (aggregated from continuous trades)"),
            ("Unit", "\u20ac/MWh"),
            ("Source", "EPEX SPOT (paid for real-time); SMARD (delayed, free)"),
            ("Cost", "Bundled with EPEX SPOT subscription"),
        ]),
        ("1.6", "Imbalance Price (reBAP) \u2014 DE", [
            ("What", "Germany\u2019s single, uniform imbalance settlement price (regelzonenuebergreifender einheitlicher Bilanzausgleichsenergiepreis)"),
            ("Why needed", "Determines cost/revenue of any deviation from contracted position; unlike UK\u2019s dual SSP/SBP, Germany uses a single symmetric price"),
            ("Frequency", "Quarter-hourly (per 15-min settlement period)"),
            ("Unit", "\u20ac/MWh"),
            ("Source", "Netztransparenz.de (published by all 4 TSOs jointly)"),
            ("Cost", "FREE"),
            ("Note", "Near-real-time estimator also available for live dashboards"),
        ]),
    ]

    for num, ttl, items in de_wholesale:
        add_data_feed_section(doc, num, ttl, items)

    # -- DE Ancillary --
    add_heading(doc, "2. Ancillary Services Clearing Prices \u2014 Germany", level=2)

    p = doc.add_paragraph(
        "Germany uses the EU-standard FCR/aFRR/mFRR framework, procured jointly by the 4 TSOs via regelleistung.net. "
        "These replace the UK\u2019s SFFR/DC/DM/DR products entirely."
    )
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True

    de_ancillary = [
        ("2.1", "FCR Capacity Price (Frequency Containment Reserve)", [
            ("What", "Clearing price for providing primary frequency containment \u2014 battery must respond within 30 seconds"),
            ("Why needed", "FCR is a high-value revenue stream for BESS \u2014 historically up to ~\u20ac9,400/MW per month"),
            ("Frequency", "Per 4-hour block (6 blocks/day), daily tender"),
            ("Unit", "\u20ac/MW per 4-hour block"),
            ("Direction", "Symmetric (single bid covers both up and down)"),
            ("Pricing", "Pay-as-cleared"),
            ("Min bid", "1 MW"),
            ("Source", "Regelleistung.net"),
            ("Cost", "FREE"),
            ("Note", "Procured via the FCR Cooperation \u2014 shared auction across DE, FR, BE, NL, AT, CH and others"),
        ]),
        ("2.2", "aFRR Positive Capacity Price (Automatic FRR \u2014 Up)", [
            ("What", "Price for providing upward secondary frequency response \u2014 activated automatically within 5 minutes"),
            ("Why needed", "aFRR is a major revenue source for German BESS \u2014 positive capacity prices have reached ~\u20ac21,500/MW per month"),
            ("Frequency", "Per 4-hour block, daily tender"),
            ("Unit", "\u20ac/MW per 4-hour block"),
            ("Pricing", "Pay-as-bid"),
            ("Min bid", "5 MW"),
            ("Source", "Regelleistung.net"),
            ("Cost", "FREE"),
        ]),
        ("2.3", "aFRR Negative Capacity Price (Automatic FRR \u2014 Down)", [
            ("What", "Price for providing downward secondary frequency response \u2014 activated when grid frequency rises"),
            ("Why needed", "Separate product from positive aFRR \u2014 enables full revenue stacking analysis"),
            ("Frequency", "Per 4-hour block, daily tender"),
            ("Unit", "\u20ac/MW per 4-hour block"),
            ("Pricing", "Pay-as-bid"),
            ("Source", "Regelleistung.net"),
            ("Cost", "FREE"),
        ]),
        ("2.4", "aFRR Energy Activation Price", [
            ("What", "Price paid when aFRR reserve is actually dispatched \u2014 separate from the capacity holding payment"),
            ("Why needed", "Determines actual revenue when battery is activated; Germany pays both capacity AND energy, unlike UK ancillary services which are capacity-only"),
            ("Frequency", "Per activation event / quarter-hourly settlement"),
            ("Unit", "\u20ac/MWh"),
            ("Pricing", "Pay-as-bid (merit order activation)"),
            ("Source", "Regelleistung.net / Netztransparenz.de"),
            ("Cost", "FREE"),
        ]),
        ("2.5", "mFRR Positive Capacity Price (Manual FRR \u2014 Up)", [
            ("What", "Price for providing upward tertiary reserve \u2014 manually activated within 15 minutes for larger imbalances"),
            ("Why needed", "Additional revenue stream, though typically lower value than FCR/aFRR for BESS"),
            ("Frequency", "Per 4-hour block, daily tender"),
            ("Unit", "\u20ac/MW per 4-hour block"),
            ("Pricing", "Pay-as-bid"),
            ("Min bid", "5 MW"),
            ("Source", "Regelleistung.net"),
            ("Cost", "FREE"),
        ]),
        ("2.6", "mFRR Negative Capacity Price (Manual FRR \u2014 Down)", [
            ("What", "Price for providing downward tertiary reserve"),
            ("Why needed", "Completes the full ancillary service product suite"),
            ("Frequency", "Per 4-hour block, daily tender"),
            ("Unit", "\u20ac/MW per 4-hour block"),
            ("Source", "Regelleistung.net"),
            ("Cost", "FREE"),
        ]),
        ("2.7", "mFRR Energy Activation Price", [
            ("What", "Price paid when mFRR is dispatched"),
            ("Why needed", "As with aFRR, actual revenue depends on both capacity and energy payments"),
            ("Frequency", "Per activation event"),
            ("Unit", "\u20ac/MWh"),
            ("Pricing", "Pay-as-bid"),
            ("Source", "Regelleistung.net / Netztransparenz.de"),
            ("Cost", "FREE"),
        ]),
    ]

    for num, ttl, items in de_ancillary:
        add_data_feed_section(doc, num, ttl, items)

    # -- DE Sources Summary --
    add_heading(doc, "3. German Data Sources Summary", level=2)

    add_table(doc,
        ["Source", "What It Provides", "Cost", "Auth Required"],
        [
            ["SMARD (Bundesnetzagentur)", "DA prices, generation mix, consumption, wholesale data", "FREE", "None"],
            ["Netztransparenz.de (4 TSOs)", "Imbalance price (reBAP), near-real-time estimator", "FREE", "None"],
            ["Regelleistung.net (4 TSOs)", "FCR/aFRR/mFRR tender results, clearing prices, activation data", "FREE", "None"],
            ["ENTSO-E Transparency Platform", "Cross-border flows, DA prices, load, generation by type", "FREE", "API token (email request)"],
            ["EPEX SPOT (direct)", "Official exchange prices, intraday continuous, order books", "Paid (quote required)", "Subscription"],
            ["Nord Pool (alternative)", "DE DA + intraday prices", "~\u20ac1,200\u20135,000/year", "Subscription"],
        ],
    )

    # ==============================
    # PART 2: NETHERLANDS
    # ==============================
    doc.add_page_break()
    add_heading(doc, "PART 2: NETHERLANDS (NL)", level=1)

    # -- NL Wholesale --
    add_heading(doc, "4. Wholesale Electricity Prices \u2014 Netherlands", level=2)

    nl_wholesale = [
        ("4.1", "EPEX Day-Ahead Auction Price (NL)", [
            ("What", "Next-day electricity price from the EPEX SPOT auction for the Netherlands bidding zone (formerly APX, now EPEX SPOT since 2015)"),
            ("Why needed", "Primary arbitrage benchmark \u2014 determines charge/discharge schedule"),
            ("Frequency", "Quarter-hourly (96 values per day, since October 2025)"),
            ("Unit", "\u20ac/MWh"),
            ("Free source", "ENTSO-E Transparency Platform"),
            ("Paid source", "EPEX SPOT (quote required)"),
            ("Cost", "FREE via ENTSO-E"),
        ]),
        ("4.2", "EPEX Intraday Auction 1 (IDA1) Price \u2014 NL", [
            ("What", "First same-day auction on EPEX SPOT for NL delivery"),
            ("Why needed", "Refines trading position after DA auction; captures updated forecasts"),
            ("Frequency", "Quarter-hourly"),
            ("Unit", "\u20ac/MWh"),
            ("Gate closure", "15:00 CET on D-1"),
            ("Source", "EPEX SPOT (paid subscription)"),
            ("Cost", "Bundled with EPEX SPOT subscription"),
        ]),
        ("4.3", "EPEX Intraday Auction 2 (IDA2) Price \u2014 NL", [
            ("What", "Second same-day auction \u2014 evening update"),
            ("Why needed", "Further position refinement"),
            ("Frequency", "Quarter-hourly"),
            ("Unit", "\u20ac/MWh"),
            ("Gate closure", "22:00 CET on D-1"),
            ("Source", "EPEX SPOT (paid subscription)"),
            ("Cost", "Bundled with EPEX SPOT subscription"),
        ]),
        ("4.4", "EPEX Intraday Auction 3 (IDA3) Price \u2014 NL", [
            ("What", "Third intraday auction \u2014 morning of delivery day"),
            ("Why needed", "Same-day demand and generation adjustments"),
            ("Frequency", "Quarter-hourly (covers 12:00\u201324:00)"),
            ("Unit", "\u20ac/MWh"),
            ("Gate closure", "10:00 CET on delivery day"),
            ("Source", "EPEX SPOT (paid subscription)"),
            ("Cost", "Bundled with EPEX SPOT subscription"),
        ]),
        ("4.5", "Intraday Continuous (IDC) Price \u2014 NL", [
            ("What", "Continuous trading price on EPEX intraday \u2014 trades up to 5 minutes before delivery"),
            ("Why needed", "Captures real-time volatility and short-notice opportunities"),
            ("Frequency", "Quarter-hourly (aggregated)"),
            ("Unit", "\u20ac/MWh"),
            ("Source", "EPEX SPOT (paid subscription for real-time)"),
            ("Cost", "Bundled with EPEX SPOT subscription"),
        ]),
        ("4.6", "Imbalance / Settlement Price \u2014 NL", [
            ("What", "TenneT\u2019s settlement price for balancing energy \u2014 determined by the most extreme aFRR activation price in each 15-minute period"),
            ("Why needed", "The Dutch market explicitly rewards \u201cpassive balancing\u201d \u2014 a BESS that charges when the system is long (or discharges when short) can profit from the imbalance price. This is a distinct revenue stream."),
            ("Frequency", "Quarter-hourly (per 15-min ISP)"),
            ("Unit", "\u20ac/MWh"),
            ("Pricing", "Usually single price; dual pricing when both up and down balancing activated in same period"),
            ("Source", "TenneT API (developer.tennet.eu)"),
            ("Cost", "FREE (registration required)"),
            ("Note", "TenneT also publishes balance delta (imbalance direction) in near-real-time \u2014 useful for passive balancing signals"),
        ]),
    ]

    for num, ttl, items in nl_wholesale:
        add_data_feed_section(doc, num, ttl, items)

    # -- NL Ancillary --
    add_heading(doc, "5. Ancillary Services Clearing Prices \u2014 Netherlands", level=2)

    p = doc.add_paragraph(
        "The Netherlands uses the same EU-standard FCR/aFRR/mFRR framework as Germany. "
        "Products are procured by TenneT NL, with cross-border sharing via PICASSO and MARI platforms."
    )
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True

    nl_ancillary = [
        ("5.1", "FCR Capacity Price \u2014 NL", [
            ("What", "Clearing price from the daily FCR Cooperation common auction"),
            ("Why needed", "FCR is a key revenue stream for Dutch BESS \u2014 procured via the same pan-European auction as Germany"),
            ("Frequency", "Per 4-hour block (6 blocks/day), daily auction"),
            ("Unit", "\u20ac/MW per 4-hour block"),
            ("Direction", "Symmetric"),
            ("Pricing", "Pay-as-cleared"),
            ("Source", "TenneT transparency data / FCR Cooperation results"),
            ("Cost", "FREE"),
            ("Note", "30% of TenneT\u2019s FCR obligation must be sourced from providers physically in the Netherlands; remaining 70% can come from any FCR Cooperation country"),
        ]),
        ("5.2", "aFRR Positive Capacity Price \u2014 NL", [
            ("What", "Price for providing upward automatic frequency restoration reserve"),
            ("Why needed", "aFRR is a significant BESS revenue source; energy activation now cross-border via PICASSO platform (since October 2024)"),
            ("Frequency", "Weekly and daily tenders"),
            ("Unit", "\u20ac/MW per period"),
            ("Source", "TenneT transparency data"),
            ("Cost", "FREE"),
        ]),
        ("5.3", "aFRR Negative Capacity Price \u2014 NL", [
            ("What", "Price for providing downward automatic frequency restoration reserve"),
            ("Why needed", "Separate product \u2014 needed for full revenue stacking analysis"),
            ("Frequency", "Weekly and daily tenders"),
            ("Unit", "\u20ac/MW per period"),
            ("Source", "TenneT transparency data"),
            ("Cost", "FREE"),
        ]),
        ("5.4", "aFRR Energy Activation Price \u2014 NL", [
            ("What", "Price paid when aFRR is dispatched via the PICASSO platform"),
            ("Why needed", "Since PICASSO go-live in NL (October 2024), cross-border aFRR sharing has created new energy activation revenue opportunities"),
            ("Frequency", "Per activation / quarter-hourly"),
            ("Unit", "\u20ac/MWh"),
            ("Source", "TenneT API / PICASSO results"),
            ("Cost", "FREE"),
        ]),
        ("5.5", "mFRR Capacity Price \u2014 NL", [
            ("What", "Price for providing manual frequency restoration reserve (historically called \u201cnoodvermogen\u201d in Dutch market)"),
            ("Why needed", "Additional revenue stream; NL connected to MARI platform from December 2025, enabling cross-border mFRR energy exchange"),
            ("Frequency", "Per tender period"),
            ("Unit", "\u20ac/MW per period"),
            ("Source", "TenneT transparency data"),
            ("Cost", "FREE"),
        ]),
        ("5.6", "mFRR Energy Activation Price \u2014 NL", [
            ("What", "Price paid when mFRR is dispatched via MARI platform"),
            ("Why needed", "New revenue stream since MARI connection in December 2025"),
            ("Frequency", "Per activation event"),
            ("Unit", "\u20ac/MWh"),
            ("Source", "TenneT API / MARI results"),
            ("Cost", "FREE"),
        ]),
    ]

    for num, ttl, items in nl_ancillary:
        add_data_feed_section(doc, num, ttl, items)

    # -- NL Sources Summary --
    add_heading(doc, "6. Dutch Data Sources Summary", level=2)

    add_table(doc,
        ["Source", "What It Provides", "Cost", "Auth Required"],
        [
            ["TenneT API (developer.tennet.eu)", "Settlement prices, imbalance volumes, balance delta, regulation state", "FREE", "Registration + token"],
            ["TenneT Transparency Portal", "Ancillary service results, system data, download CSVs", "FREE", "Token for automation"],
            ["ENTSO-E Transparency Platform", "DA prices, cross-border flows, generation, load", "FREE", "API token (email request)"],
            ["EPEX SPOT (direct)", "Official NL exchange prices, intraday continuous, order books", "Paid (quote required)", "Subscription"],
            ["Nord Pool (alternative)", "NL DA + intraday prices", "~\u20ac1,700\u20133,300/year (FR/NL/AT/BE bundle)", "Subscription"],
        ],
    )

    # ==============================
    # PART 3: CROSS-MARKET
    # ==============================
    doc.add_page_break()
    add_heading(doc, "PART 3: CROSS-MARKET COMPARISON", level=1)

    # Feed count
    add_heading(doc, "7. Data Feed Count by Market", level=2)

    add_table(doc,
        ["Category", "UK (GB)", "Germany (DE)", "Netherlands (NL)"],
        [
            ["Wholesale prices", "7 feeds", "6 feeds", "6 feeds"],
            ["Imbalance prices", "2 (SSP + SBP)", "1 (reBAP)", "1 (settlement price)"],
            ["Ancillary capacity prices", "7 (SFFR, DCL/H, DML/H, DRL/H)", "5 (FCR, aFRR+/-, mFRR+/-)", "5 (FCR, aFRR+/-, mFRR+/-)"],
            ["Ancillary energy prices", "0 (capacity-only)", "2 (aFRR + mFRR energy)", "2 (aFRR + mFRR energy)"],
            ["Total individual feeds", "16", "14", "14"],
        ],
    )

    # Ancillary mapping
    add_heading(doc, "8. Ancillary Services Mapping (UK vs DE/NL)", level=2)

    add_table(doc,
        ["UK Product", "DE/NL Equivalent", "Key Difference"],
        [
            ["SFFR (Static Firm Frequency Response)", "No direct equivalent", "UK-specific product; closest is FCR but architecture differs"],
            ["DC Low / DC High (Dynamic Containment)", "FCR (Frequency Containment Reserve)", "FCR is symmetric (one bid = both directions); DC is split Low/High. FCR 30s response; DC 1s"],
            ["DM Low / DM High (Dynamic Modulation)", "aFRR (automatic FRR)", "aFRR activates within 5 min; separate +/- tenders. DE/NL pay capacity AND energy; UK capacity only"],
            ["DR Low / DR High (Dynamic Regulation)", "mFRR (manual FRR)", "mFRR activates within 15 min; separate +/-. Connected to MARI cross-border platform"],
            ["Procured by NESO, EFA blocks", "Regelleistung.net (DE) / TenneT (NL), 4-hour blocks", "DE/NL: pay-as-bid for aFRR/mFRR; UK: mostly pay-as-cleared"],
        ],
    )

    # Cost summary
    add_heading(doc, "9. Full Cost Summary \u2014 All Three Markets", level=2)

    add_heading(doc, "Free Data Sources (sufficient for analytical dashboard)", level=3)

    add_table(doc,
        ["Market", "Source", "Data Provided", "Cost"],
        [
            ["UK", "Elexon BMRS", "SSP, SBP, DA HH reference", "FREE"],
            ["UK", "NESO Data Portal", "Ancillary clearing prices (all 7 services)", "FREE"],
            ["UK", "GridBeyond report", "All wholesale + ancillary (bundled in aggregator fee)", "FREE"],
            ["DE", "SMARD", "DA prices, generation mix, wholesale data", "FREE"],
            ["DE", "Netztransparenz.de", "Imbalance price (reBAP)", "FREE"],
            ["DE", "Regelleistung.net", "FCR/aFRR/mFRR tender results", "FREE"],
            ["DE", "ENTSO-E", "Cross-border flows, load, generation", "FREE"],
            ["NL", "TenneT API", "Settlement prices, imbalance data, ancillary results", "FREE"],
            ["NL", "ENTSO-E", "DA prices, cross-border flows", "FREE"],
        ],
    )

    add_heading(doc, "Paid Subscriptions (for real-time / direct feeds)", level=3)

    add_table(doc,
        ["Subscription", "Markets Covered", "Annual Cost (approx.)"],
        [
            ["EPEX SPOT DE data package", "Germany DA + Intraday", "Quote required"],
            ["EPEX SPOT NL data package", "Netherlands DA + Intraday", "Quote required"],
            ["Nord Pool DE Day-Ahead", "Germany DA prices", "~\u20ac1,200/year"],
            ["Nord Pool DE Intraday", "Germany intraday trades", "~\u20ac5,000/year"],
            ["Nord Pool NL bundle (FR/NL/AT/BE)", "Netherlands DA prices", "~\u20ac1,700\u20133,300/year"],
            ["Nord Pool Intraday (per market)", "NL intraday trades", "~\u20ac5,000/year"],
            ["Modo Energy (benchmarking)", "GB, DE, NL analytics", "~\u20ac8,000\u201360,000/year"],
        ],
    )

    add_heading(doc, "Cost Scenarios", level=3)

    add_table(doc,
        ["Scenario", "UK Cost", "DE Cost", "NL Cost", "Total (approx.)"],
        [
            ["Minimum (free sources only)", "\u00a30", "\u20ac0", "\u20ac0", "\u00a30 / \u20ac0"],
            ["Standard (with aggregator)", "\u00a30 (GridBeyond)", "Depends on aggregator", "Depends on aggregator", "Varies"],
            ["With direct market feeds", "~\u00a35,300", "~\u20ac6,200", "~\u20ac6,700\u20138,300", "~\u00a318,000\u201320,000 equiv."],
            ["Full (feeds + benchmarking)", "~\u00a320,000", "~\u20ac14,000\u201366,000", "~\u20ac14,000\u201368,000", "~\u00a340,000\u2013150,000 equiv."],
        ],
    )

    # Implementation notes
    add_heading(doc, "10. Dashboard Implementation Notes", level=2)

    add_heading(doc, "Data Model Changes Required", level=3)
    changes = [
        "Settlement period: Must support both 30-min (UK) and 15-min (DE/NL) \u2014 48 vs 96 periods/day",
        "Currency: GBP for UK, EUR for DE/NL \u2014 dashboard needs multi-currency support or conversion",
        "Ancillary module redesign: UK\u2019s 7 frequency response products do not map 1:1 to DE/NL\u2019s 5 products. The aFRR/mFRR energy activation payment does not exist in the UK model",
        "Imbalance module: UK dual pricing (SSP/SBP) vs DE single pricing (reBAP) vs NL hybrid \u2014 different logic per market",
        "Intraday auctions: UK has 1 IDA; DE/NL have 3 IDAs each",
    ]
    for c in changes:
        p = doc.add_paragraph(style="List Number")
        p.text = c
        for run in p.runs:
            run.font.size = Pt(10)

    add_heading(doc, "New Data Fields for DE/NL (not in current UK dashboard)", level=3)
    new_fields = [
        "aFRR energy activation price (\u20ac/MWh)",
        "mFRR energy activation price (\u20ac/MWh)",
        "Balance delta / system imbalance direction (for NL passive balancing)",
        "Cross-border flow data (significant price driver for NL due to high interconnection)",
        "Grid fees / network charges (significant cost line in NL; less impactful in UK/DE)",
    ]
    for f in new_fields:
        p = doc.add_paragraph(style="List Bullet")
        p.text = f
        for run in p.runs:
            run.font.size = Pt(10)

    add_heading(doc, "API Integration Priority", level=3)
    apis = [
        "SMARD (DE) \u2014 REST API, no auth, well-documented, Python library available",
        "TenneT API (NL) \u2014 REST API, token auth, good documentation",
        "ENTSO-E (both) \u2014 REST API, email for token, covers both markets",
        "Regelleistung.net (DE) \u2014 web data centre, free, no auth",
        "EPEX SPOT (both) \u2014 paid, for real-time feeds only",
    ]
    for i, a in enumerate(apis, 1):
        p = doc.add_paragraph(style="List Number")
        p.text = a
        for run in p.runs:
            run.font.size = Pt(10)

    doc.save(r"C:\repos\bess-dashboard\docs\Market_Data_Subscriptions_DE_NL.docx")
    print("DE/NL document saved.")


if __name__ == "__main__":
    generate_uk_doc()
    generate_de_nl_doc()
    print("Both documents generated successfully.")
