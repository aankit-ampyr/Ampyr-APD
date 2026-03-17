"""
Generate a single Word document with 3 market data procurement tables.
All feeds listed with paid/commercial-grade sources only.
No pricing included.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shading = tc_pr.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex},
    )
    tc_pr.append(shading)


def style_header_row(table, color_hex="1F3864"):
    for cell in table.rows[0].cells:
        set_cell_shading(cell, color_hex)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)


def add_table(doc, headers, rows, font_size=Pt(8)):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = font_size

    style_header_row(tbl)

    for r_idx in range(1, len(tbl.rows)):
        if r_idx % 2 == 0:
            for cell in tbl.rows[r_idx].cells:
                set_cell_shading(cell, "E8EDF3")

    return tbl


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    return h


# ---------------------------------------------------------------------------
# Column definitions
# ---------------------------------------------------------------------------

HEADERS = [
    "#",
    "Data Feed",
    "Category",
    "Frequency / Resolution",
    "Commercial Data Provider(s)",
    "Delivery Method",
    "SLA / Reliability Notes",
]

# ---------------------------------------------------------------------------
# UK DATA
# ---------------------------------------------------------------------------

UK_ROWS = [
    [
        "1",
        "EPEX Day-Ahead Price (GB, 60-min auction)",
        "Wholesale",
        "Half-hourly (48/day)\nPublished daily after 09:20 GMT auction",
        "EPEX SPOT\nNord Pool (N2EX)\nModo Energy",
        "EPEX: sFTP (CSV) or API\nNord Pool: REST API\nModo: REST API",
        "EPEX/Nord Pool: exchange-grade, contractual SLA\nModo: aggregated from exchange, SLA per subscription tier",
    ],
    [
        "2",
        "EPEX Day-Ahead Price (GB, 30-min auction)",
        "Wholesale",
        "Half-hourly (48/day)\nPublished daily after 15:30 GMT auction",
        "EPEX SPOT\nModo Energy",
        "EPEX: sFTP (CSV) or API\nModo: REST API",
        "UK-specific 30-min product. Not available via Nord Pool",
    ],
    [
        "3",
        "GB-ISEM Intraday Auction 1 (IDA1) Price",
        "Wholesale",
        "Half-hourly (48/day)\nPublished daily after 17:30 UTC auction",
        "EPEX SPOT\nModo Energy",
        "EPEX: sFTP (CSV) or API\nModo: REST API",
        "Coupled with SEM (Ireland) via Moyle and East-West interconnectors",
    ],
    [
        "4",
        "Intraday Continuous (IDC) Price",
        "Wholesale",
        "Half-hourly (aggregated)\nContinuous trading up to gate closure",
        "EPEX SPOT\nModo Energy",
        "EPEX: sFTP (CSV, end-of-day) or API (real-time)\nModo: REST API (near real-time)",
        "Real-time feed requires EPEX API subscription.\nEnd-of-day aggregates via sFTP",
    ],
    [
        "5",
        "DA HH Reference Price (Market Index / MID)",
        "Wholesale",
        "Half-hourly (48/day)\nPublished daily",
        "EPEX SPOT (source exchange)\nNord Pool (source exchange)\nModo Energy",
        "EPEX/Nord Pool: sFTP or API\nModo: REST API",
        "Note: MID is a composite index from both EPEX and Nord Pool short-term trades. Not identical to raw EPEX DA auction price",
    ],
    [
        "6",
        "System Sell Price (SSP)",
        "Imbalance",
        "Half-hourly (48/day)\nPublished per settlement period",
        "EPEX SPOT (bundled in GB data package)\nModo Energy\nNord Pool",
        "EPEX: sFTP or API\nModo: REST API\nNord Pool: API",
        "Authoritative source: Elexon (settlement agent). Commercial providers source from Elexon",
    ],
    [
        "7",
        "System Buy Price (SBP)",
        "Imbalance",
        "Half-hourly (48/day)\nPublished per settlement period",
        "EPEX SPOT (bundled in GB data package)\nModo Energy\nNord Pool",
        "EPEX: sFTP or API\nModo: REST API\nNord Pool: API",
        "Authoritative source: Elexon (settlement agent). Commercial providers source from Elexon",
    ],
    [
        "8",
        "SFFR (Static Firm Frequency Response)\nAvailability, Clearing Price, Revenue",
        "Ancillary Services",
        "Per EFA block (6 blocks/day)\nPublished after each tender",
        "Modo Energy\nNord Pool (GB ancillary package)",
        "Modo: REST API\nNord Pool: API",
        "Authoritative source: NESO. Commercial providers source from NESO tender results",
    ],
    [
        "9",
        "DCL (Dynamic Containment Low)\nAvailability, Clearing Price, Revenue",
        "Ancillary Services",
        "Per EFA block (6 blocks/day)",
        "Modo Energy\nNord Pool (GB ancillary package)",
        "Modo: REST API\nNord Pool: API",
        "Part of DC/DM/DR product suite procured by NESO",
    ],
    [
        "10",
        "DCH (Dynamic Containment High)\nAvailability, Clearing Price, Revenue",
        "Ancillary Services",
        "Per EFA block (6 blocks/day)",
        "Modo Energy\nNord Pool (GB ancillary package)",
        "Modo: REST API\nNord Pool: API",
        "",
    ],
    [
        "11",
        "DML (Dynamic Modulation Low)\nAvailability, Clearing Price, Revenue",
        "Ancillary Services",
        "Per EFA block (6 blocks/day)",
        "Modo Energy\nNord Pool (GB ancillary package)",
        "Modo: REST API\nNord Pool: API",
        "",
    ],
    [
        "12",
        "DMH (Dynamic Modulation High)\nAvailability, Clearing Price, Revenue",
        "Ancillary Services",
        "Per EFA block (6 blocks/day)",
        "Modo Energy\nNord Pool (GB ancillary package)",
        "Modo: REST API\nNord Pool: API",
        "",
    ],
    [
        "13",
        "DRL (Dynamic Regulation Low)\nAvailability, Clearing Price, Revenue",
        "Ancillary Services",
        "Per EFA block (6 blocks/day)",
        "Modo Energy\nNord Pool (GB ancillary package)",
        "Modo: REST API\nNord Pool: API",
        "",
    ],
    [
        "14",
        "DRH (Dynamic Regulation High)\nAvailability, Clearing Price, Revenue",
        "Ancillary Services",
        "Per EFA block (6 blocks/day)",
        "Modo Energy\nNord Pool (GB ancillary package)",
        "Modo: REST API\nNord Pool: API",
        "",
    ],
    [
        "15",
        "Imbalance Settlement Data\n(Revenue, Charges, Volumes)",
        "Imbalance",
        "Half-hourly (48/day)",
        "EPEX SPOT (bundled)\nModo Energy\nNord Pool",
        "EPEX: sFTP or API\nModo: REST API\nNord Pool: API",
        "Authoritative source: Elexon settlement",
    ],
]

# ---------------------------------------------------------------------------
# GERMANY DATA
# ---------------------------------------------------------------------------

DE_ROWS = [
    [
        "1",
        "EPEX Day-Ahead Price (DE-LU, 15-min)",
        "Wholesale",
        "Quarter-hourly (96/day)\nPublished daily after 12:00 CET auction",
        "EPEX SPOT\nNord Pool\nModo Energy",
        "EPEX: sFTP (CSV) or API\nNord Pool: REST API\nModo: REST API",
        "15-min granularity since October 2025 (previously hourly). Bidding zone: DE-LU",
    ],
    [
        "2",
        "EPEX Intraday Auction 1 (IDA1) Price \u2014 DE",
        "Wholesale",
        "Quarter-hourly (96/day)\nGate closure: 15:00 CET D-1",
        "EPEX SPOT\nModo Energy",
        "EPEX: sFTP (CSV) or API\nModo: REST API",
        "First of 3 intraday auctions",
    ],
    [
        "3",
        "EPEX Intraday Auction 2 (IDA2) Price \u2014 DE",
        "Wholesale",
        "Quarter-hourly (96/day)\nGate closure: 22:00 CET D-1",
        "EPEX SPOT\nModo Energy",
        "EPEX: sFTP (CSV) or API\nModo: REST API",
        "",
    ],
    [
        "4",
        "EPEX Intraday Auction 3 (IDA3) Price \u2014 DE",
        "Wholesale",
        "Quarter-hourly (covers 12:00\u201324:00)\nGate closure: 10:00 CET delivery day",
        "EPEX SPOT\nModo Energy",
        "EPEX: sFTP (CSV) or API\nModo: REST API",
        "",
    ],
    [
        "5",
        "Intraday Continuous (IDC) Price \u2014 DE",
        "Wholesale",
        "Quarter-hourly (aggregated)\nContinuous trading, gate closure 5 min before delivery",
        "EPEX SPOT\nModo Energy",
        "EPEX: sFTP (end-of-day) or API (real-time)\nModo: REST API",
        "5-minute gate closure (vs UK 30-min). Real-time feed via EPEX API",
    ],
    [
        "6",
        "Imbalance Price (reBAP)",
        "Imbalance",
        "Quarter-hourly (96/day)\nPublished per 15-min settlement period",
        "EPEX SPOT (bundled in DE package)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "Germany uses single symmetric imbalance price (unlike UK dual SSP/SBP). Authoritative source: Netztransparenz.de (4 TSOs)",
    ],
    [
        "7",
        "FCR Capacity Price\n(Frequency Containment Reserve)",
        "Ancillary Services",
        "Per 4-hour block (6 blocks/day)\nDaily tender",
        "EPEX SPOT (bundled)\nModo Energy\nNord Pool",
        "EPEX: sFTP or API\nModo: REST API\nNord Pool: API",
        "Symmetric product, pay-as-cleared. Same FCR Cooperation auction as NL. Authoritative source: Regelleistung.net",
    ],
    [
        "8",
        "aFRR Positive Capacity Price\n(Automatic FRR \u2014 Up)",
        "Ancillary Services",
        "Per 4-hour block (6 blocks/day)\nDaily tender",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "Pay-as-bid. Min bid: 5 MW. Authoritative source: Regelleistung.net",
    ],
    [
        "9",
        "aFRR Negative Capacity Price\n(Automatic FRR \u2014 Down)",
        "Ancillary Services",
        "Per 4-hour block (6 blocks/day)\nDaily tender",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "Separate product from aFRR positive",
    ],
    [
        "10",
        "aFRR Energy Activation Price",
        "Ancillary Services",
        "Per activation / quarter-hourly settlement",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "Energy payment on top of capacity \u2014 concept not present in UK model. Via PICASSO platform",
    ],
    [
        "11",
        "mFRR Positive Capacity Price\n(Manual FRR \u2014 Up)",
        "Ancillary Services",
        "Per 4-hour block (6 blocks/day)\nDaily tender",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "Pay-as-bid. Min bid: 5 MW. Authoritative source: Regelleistung.net",
    ],
    [
        "12",
        "mFRR Negative Capacity Price\n(Manual FRR \u2014 Down)",
        "Ancillary Services",
        "Per 4-hour block (6 blocks/day)\nDaily tender",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "",
    ],
    [
        "13",
        "mFRR Energy Activation Price",
        "Ancillary Services",
        "Per activation event",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "Via MARI platform. Energy payment on top of capacity",
    ],
]

# ---------------------------------------------------------------------------
# NETHERLANDS DATA
# ---------------------------------------------------------------------------

NL_ROWS = [
    [
        "1",
        "EPEX Day-Ahead Price (NL, 15-min)",
        "Wholesale",
        "Quarter-hourly (96/day)\nPublished daily after 12:00 CET auction",
        "EPEX SPOT\nNord Pool\nModo Energy",
        "EPEX: sFTP (CSV) or API\nNord Pool: REST API (FR/NL/AT/BE bundle)\nModo: REST API",
        "15-min granularity since October 2025. Bidding zone: NL (formerly APX, now EPEX SPOT since 2015)",
    ],
    [
        "2",
        "EPEX Intraday Auction 1 (IDA1) Price \u2014 NL",
        "Wholesale",
        "Quarter-hourly (96/day)\nGate closure: 15:00 CET D-1",
        "EPEX SPOT\nModo Energy",
        "EPEX: sFTP (CSV) or API\nModo: REST API",
        "",
    ],
    [
        "3",
        "EPEX Intraday Auction 2 (IDA2) Price \u2014 NL",
        "Wholesale",
        "Quarter-hourly (96/day)\nGate closure: 22:00 CET D-1",
        "EPEX SPOT\nModo Energy",
        "EPEX: sFTP (CSV) or API\nModo: REST API",
        "",
    ],
    [
        "4",
        "EPEX Intraday Auction 3 (IDA3) Price \u2014 NL",
        "Wholesale",
        "Quarter-hourly (covers 12:00\u201324:00)\nGate closure: 10:00 CET delivery day",
        "EPEX SPOT\nModo Energy",
        "EPEX: sFTP (CSV) or API\nModo: REST API",
        "",
    ],
    [
        "5",
        "Intraday Continuous (IDC) Price \u2014 NL",
        "Wholesale",
        "Quarter-hourly (aggregated)\nContinuous trading, gate closure 5 min before delivery",
        "EPEX SPOT\nModo Energy",
        "EPEX: sFTP (end-of-day) or API (real-time)\nModo: REST API",
        "5-minute gate closure",
    ],
    [
        "6",
        "Imbalance / Settlement Price \u2014 NL",
        "Imbalance",
        "Quarter-hourly (96/day)\nPer 15-min ISP",
        "EPEX SPOT (bundled in NL package)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "Hybrid single/dual pricing. Determined by most extreme aFRR activation price. Authoritative source: TenneT NL",
    ],
    [
        "7",
        "Balance Delta (System Imbalance Direction) \u2014 NL",
        "Imbalance",
        "Near real-time (12-second snapshots since Feb 2026)",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "NL-specific: used for passive balancing strategy. Authoritative source: TenneT NL",
    ],
    [
        "8",
        "FCR Capacity Price \u2014 NL",
        "Ancillary Services",
        "Per 4-hour block (6 blocks/day)\nDaily auction",
        "EPEX SPOT (bundled)\nModo Energy\nNord Pool",
        "EPEX: sFTP or API\nModo: REST API\nNord Pool: API",
        "Same FCR Cooperation auction as DE. Symmetric, pay-as-cleared. 30% must be from NL-connected providers",
    ],
    [
        "9",
        "aFRR Positive Capacity Price \u2014 NL",
        "Ancillary Services",
        "Weekly and daily tenders",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "Authoritative source: TenneT NL. Energy activation via PICASSO platform (since Oct 2024)",
    ],
    [
        "10",
        "aFRR Negative Capacity Price \u2014 NL",
        "Ancillary Services",
        "Weekly and daily tenders",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "",
    ],
    [
        "11",
        "aFRR Energy Activation Price \u2014 NL",
        "Ancillary Services",
        "Per activation / quarter-hourly",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "Via PICASSO platform. Cross-border aFRR sharing since Oct 2024",
    ],
    [
        "12",
        "mFRR Capacity Price \u2014 NL",
        "Ancillary Services",
        "Per tender period",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "Historically called 'noodvermogen'. Authoritative source: TenneT NL",
    ],
    [
        "13",
        "mFRR Energy Activation Price \u2014 NL",
        "Ancillary Services",
        "Per activation event",
        "EPEX SPOT (bundled)\nModo Energy",
        "EPEX: sFTP or API\nModo: REST API",
        "Via MARI platform since Dec 2025",
    ],
]


# ---------------------------------------------------------------------------
# Generate document
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # === Title ===
    title = doc.add_heading("Market Data Requirements \u2014 Commercial Subscriptions", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    meta = doc.add_paragraph()
    for lbl, val in [
        ("Purpose", "Procurement of commercial-grade market data subscriptions with API access for automated ingestion"),
        ("Markets", "United Kingdom (GB), Germany (DE), Netherlands (NL)"),
        ("Requirement", "Reliable API/sFTP access with contractual SLA for continuous data delivery. "
                        "Data will be stored historically for ongoing analysis."),
        ("Prepared", "February 2026"),
    ]:
        r = meta.add_run(f"{lbl}: ")
        r.bold = True
        r.font.size = Pt(10)
        r2 = meta.add_run(f"{val}\n")
        r2.font.size = Pt(10)

    doc.add_paragraph()

    # Key context
    add_heading(doc, "Key Requirements", level=2)
    reqs = [
        "All data feeds must be accessible via API or automated sFTP \u2014 no manual downloads",
        "Commercial licence permitting internal use, storage, and integration into proprietary applications",
        "Contractual SLA for uptime and data delivery timeliness",
        "Historical data backfill capability (minimum 12 months)",
        "Consistent data format and schema across all feeds within a provider",
    ]
    for req in reqs:
        p = doc.add_paragraph(style="List Bullet")
        p.text = req
        for run in p.runs:
            run.font.size = Pt(9)

    doc.add_paragraph()

    # =====================================================================
    # TABLE 1: UK
    # =====================================================================
    add_heading(doc, "Table 1: United Kingdom (GB)", level=1)

    p = doc.add_paragraph()
    specs = [
        ("Settlement period", "30 minutes (48 periods/day)"),
        ("Currency", "GBP"),
        ("TSO", "NESO"),
        ("Exchange", "EPEX SPOT / Nord Pool"),
    ]
    for i, (lbl, val) in enumerate(specs):
        r = p.add_run(f"{lbl}: ")
        r.bold = True
        r.font.size = Pt(9)
        sep = "  |  " if i < len(specs) - 1 else ""
        r2 = p.add_run(f"{val}{sep}")
        r2.font.size = Pt(9)

    add_table(doc, HEADERS, UK_ROWS)
    doc.add_paragraph()

    # =====================================================================
    # TABLE 2: GERMANY
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, "Table 2: Germany (DE)", level=1)

    p = doc.add_paragraph()
    specs = [
        ("Settlement period", "15 minutes (96 periods/day)"),
        ("Currency", "EUR"),
        ("TSOs", "50Hertz, Amprion, TenneT, TransnetBW"),
        ("Exchange", "EPEX SPOT"),
    ]
    for i, (lbl, val) in enumerate(specs):
        r = p.add_run(f"{lbl}: ")
        r.bold = True
        r.font.size = Pt(9)
        sep = "  |  " if i < len(specs) - 1 else ""
        r2 = p.add_run(f"{val}{sep}")
        r2.font.size = Pt(9)

    add_table(doc, HEADERS, DE_ROWS)
    doc.add_paragraph()

    # =====================================================================
    # TABLE 3: NETHERLANDS
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, "Table 3: Netherlands (NL)", level=1)

    p = doc.add_paragraph()
    specs = [
        ("Settlement period", "15 minutes (96 periods/day)"),
        ("Currency", "EUR"),
        ("TSO", "TenneT NL"),
        ("Exchange", "EPEX SPOT (formerly APX)"),
    ]
    for i, (lbl, val) in enumerate(specs):
        r = p.add_run(f"{lbl}: ")
        r.bold = True
        r.font.size = Pt(9)
        sep = "  |  " if i < len(specs) - 1 else ""
        r2 = p.add_run(f"{val}{sep}")
        r2.font.size = Pt(9)

    add_table(doc, HEADERS, NL_ROWS)
    doc.add_paragraph()

    # =====================================================================
    # COMMERCIAL PROVIDERS SUMMARY
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, "Commercial Data Providers \u2014 Summary", level=1)

    p = doc.add_paragraph(
        "The following providers can supply the data listed above. "
        "Please advise which providers we have existing relationships with and recommend the optimal procurement approach."
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].italic = True

    provider_headers = [
        "Provider",
        "Markets Covered",
        "Data Categories",
        "Delivery Methods",
        "Contact",
    ]

    provider_rows = [
        [
            "EPEX SPOT\n(EEX Group)",
            "UK (GB), Germany (DE-LU), Netherlands (NL)",
            "All wholesale prices (DA, IDA, IDC)\nImbalance prices\nAncillary services data",
            "sFTP (CSV files, daily)\nAPI (real-time or near real-time)",
            "marketdata.sales@epexspot.com\nepexspot.com/en/marketdataservices",
        ],
        [
            "Nord Pool",
            "UK (N2EX DA), Germany (DE DA), Netherlands (NL DA)\nIntraday (selected markets)",
            "Day-ahead prices\nIntraday trade data\nAncillary services (selected)",
            "REST API\n1 API account included per subscription",
            "nordpoolgroup.com/en/services/power-market-data-services",
        ],
        [
            "Modo Energy",
            "UK (primary), Germany, Netherlands (expanding coverage)",
            "All EPEX wholesale prices\nBESS revenue benchmarks\nAncillary services\nImbalance data",
            "REST API\nMultiple subscription tiers available",
            "modoenergy.com/pricing\ndevelopers.modoenergy.com",
        ],
    ]

    add_table(doc, provider_headers, provider_rows)

    doc.add_paragraph()

    # Action required
    add_heading(doc, "Action Required", level=2)
    actions = [
        "Confirm whether we have any existing subscriptions or contracts with the above providers",
        "For each data feed in the tables above, confirm which provider and subscription tier is recommended",
        "Confirm API access is included (not just sFTP/manual download)",
        "Confirm historical data backfill is available (minimum 12 months) for each feed",
        "Provide subscription options with coverage across all three markets (UK, DE, NL) where possible to minimise the number of vendor contracts",
    ]
    for i, a in enumerate(actions, 1):
        p = doc.add_paragraph(style="List Number")
        p.text = a
        for run in p.runs:
            run.font.size = Pt(9)

    # Save
    output_path = r"C:\repos\bess-dashboard\docs\Market_Data_Procurement.docx"
    doc.save(output_path)
    print(f"Document saved: {output_path}")


if __name__ == "__main__":
    main()
