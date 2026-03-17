"""
Generate a single Word document with three market data subscription tables:
  Table 1: UK (GB)
  Table 2: Germany (DE)
  Table 3: Netherlands (NL)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Utility helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shading = tc_pr.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex},
    )
    tc_pr.append(shading)


def style_header_row(table, color_hex="1F3864"):
    for cell in table.rows[0].cells:
        set_cell_shading(cell, color_hex)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)


def add_table(doc, headers, rows, font_size=Pt(8)):
    """Insert a formatted table and return it."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = font_size

    style_header_row(tbl)

    # Alternating row shading
    for r_idx in range(1, len(tbl.rows)):
        if r_idx % 2 == 0:
            for cell in tbl.rows[r_idx].cells:
                set_cell_shading(cell, "E8EDF3")

    return tbl


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    return h


# ---------------------------------------------------------------------------
# Table data
# ---------------------------------------------------------------------------

UK_HEADERS = [
    "#",
    "Data Field",
    "In GridBeyond Report",
    "Automated API Source (Free)",
    "Automated API Source (Paid)",
    "Subscription Tiers & Cost",
]

UK_ROWS = [
    [
        "1",
        "EPEX Day-Ahead Price (60-min)",
        "Yes (HH)",
        "None \u2014 Elexon MID is a composite proxy, not the actual EPEX DA auction price",
        "EPEX SPOT sFTP\nModo Energy API\nNord Pool (N2EX)",
        "EPEX SPOT: quote required\nNord Pool N2EX DA: ~\u20ac1,200/yr\nModo Data Pro: ~\u00a38,000\u201315,000/yr",
    ],
    [
        "2",
        "EPEX 30-Min Day-Ahead Price",
        "Yes (HH)",
        "None",
        "EPEX SPOT sFTP\nModo Energy API",
        "EPEX SPOT: quote required\nModo Data Pro: ~\u00a38,000\u201315,000/yr",
    ],
    [
        "3",
        "GB-ISEM IDA1 Price",
        "Yes (HH)",
        "None",
        "EPEX SPOT sFTP\nModo Energy API",
        "EPEX SPOT: quote required\nModo Data Pro: ~\u00a38,000\u201315,000/yr",
    ],
    [
        "4",
        "Intraday Continuous (IDC) Price",
        "Yes (HH)",
        "None",
        "EPEX SPOT sFTP\nModo Energy API",
        "EPEX SPOT: quote required\nModo Data Pro: ~\u00a38,000\u201315,000/yr",
    ],
    [
        "5",
        "DA HH Reference Price (Market Index)",
        "Yes (HH)",
        "Elexon BMRS API \u2014 no auth, 5,000 req/min\n(Note: composite index, not raw EPEX price)",
        "\u2014",
        "FREE",
    ],
    [
        "6",
        "System Sell Price (SSP)",
        "Yes (HH)",
        "Elexon BMRS API \u2014 no auth, 5,000 req/min",
        "\u2014",
        "FREE",
    ],
    [
        "7",
        "System Buy Price (SBP)",
        "Yes (HH)",
        "Elexon BMRS API \u2014 no auth, 5,000 req/min",
        "\u2014",
        "FREE",
    ],
    [
        "8",
        "SFFR (Availability, Price, Revenue)",
        "Yes (per EFA block)",
        "NESO Data Portal API (CKAN) \u2014 no auth, 2 req/min",
        "\u2014",
        "FREE",
    ],
    [
        "9",
        "DCL (Availability, Price, Revenue)",
        "Yes (per EFA block)",
        "NESO Data Portal API (CKAN) \u2014 no auth, 2 req/min",
        "\u2014",
        "FREE",
    ],
    [
        "10",
        "DCH (Availability, Price, Revenue)",
        "Yes (per EFA block)",
        "NESO Data Portal API (CKAN) \u2014 no auth, 2 req/min",
        "\u2014",
        "FREE",
    ],
    [
        "11",
        "DML (Availability, Price, Revenue)",
        "Yes (per EFA block)",
        "NESO Data Portal API (CKAN) \u2014 no auth, 2 req/min",
        "\u2014",
        "FREE",
    ],
    [
        "12",
        "DMH (Availability, Price, Revenue)",
        "Yes (per EFA block)",
        "NESO Data Portal API (CKAN) \u2014 no auth, 2 req/min",
        "\u2014",
        "FREE",
    ],
    [
        "13",
        "DRL (Availability, Price, Revenue)",
        "Yes (per EFA block)",
        "NESO Data Portal API (CKAN) \u2014 no auth, 2 req/min",
        "\u2014",
        "FREE",
    ],
    [
        "14",
        "DRH (Availability, Price, Revenue)",
        "Yes (per EFA block)",
        "NESO Data Portal API (CKAN) \u2014 no auth, 2 req/min",
        "\u2014",
        "FREE",
    ],
    [
        "15",
        "Imbalance Revenue & Charges",
        "Yes (HH)",
        "Elexon BMRS API \u2014 no auth, 5,000 req/min",
        "\u2014",
        "FREE",
    ],
]

# ---------------------------------------------------------------------------

DE_HEADERS = [
    "#",
    "Data Field",
    "Automated API Source (Free)",
    "Automated API Source (Paid)",
    "Subscription Tiers & Cost",
]

DE_ROWS = [
    [
        "1",
        "EPEX Day-Ahead Price (DE-LU, 15-min)",
        "SMARD API \u2014 no auth, JSON\nENTSO-E API \u2014 free token\nEnergy-Charts API \u2014 no auth",
        "\u2014",
        "FREE (3 independent sources)",
    ],
    [
        "2",
        "EPEX Intraday Auction 1 (IDA1) Price",
        "None (partial data on SMARD, not clean IDA1 series)",
        "EPEX SPOT sFTP\nModo Energy API",
        "EPEX SPOT: quote required\nModo Data Pro: ~\u20ac8,000\u201315,000/yr",
    ],
    [
        "3",
        "EPEX Intraday Auction 2 (IDA2) Price",
        "None",
        "EPEX SPOT sFTP\nModo Energy API",
        "EPEX SPOT: quote required\nModo Data Pro: ~\u20ac8,000\u201315,000/yr",
    ],
    [
        "4",
        "EPEX Intraday Auction 3 (IDA3) Price",
        "None",
        "EPEX SPOT sFTP\nModo Energy API",
        "EPEX SPOT: quote required\nModo Data Pro: ~\u20ac8,000\u201315,000/yr",
    ],
    [
        "5",
        "Intraday Continuous (IDC) Price",
        "SMARD API (delayed/aggregated only)",
        "EPEX SPOT sFTP (real-time)\nModo Energy API",
        "Delayed: FREE via SMARD\nReal-time: EPEX quote required / Modo ~\u20ac8,000\u201315,000/yr",
    ],
    [
        "6",
        "Imbalance Price (reBAP)",
        "Netztransparenz.de WebAPI \u2014 OAuth 2.0, free registration\nENTSO-E API \u2014 free token",
        "\u2014",
        "FREE (2 sources; Netztransparenz is authoritative)",
    ],
    [
        "7",
        "FCR Capacity Price",
        "Regelleistung.net \u2014 HTTP/XLSX download, no auth (undocumented endpoint)\nENTSO-E API \u2014 free token",
        "\u2014",
        "FREE (note: Regelleistung endpoint is undocumented, may change)",
    ],
    [
        "8",
        "aFRR Positive Capacity Price",
        "Regelleistung.net \u2014 HTTP/XLSX, no auth\nENTSO-E API \u2014 free token",
        "\u2014",
        "FREE",
    ],
    [
        "9",
        "aFRR Negative Capacity Price",
        "Regelleistung.net \u2014 HTTP/XLSX, no auth\nENTSO-E API \u2014 free token",
        "\u2014",
        "FREE",
    ],
    [
        "10",
        "aFRR Energy Activation Price",
        "ENTSO-E API (A84 doc type) \u2014 free token\nRegelleistung.net \u2014 XLSX\nNetztransparenz.de \u2014 OAuth",
        "\u2014",
        "FREE (3 sources)",
    ],
    [
        "11",
        "mFRR Positive Capacity Price",
        "Regelleistung.net \u2014 HTTP/XLSX, no auth\nENTSO-E API \u2014 free token",
        "\u2014",
        "FREE",
    ],
    [
        "12",
        "mFRR Negative Capacity Price",
        "Regelleistung.net \u2014 HTTP/XLSX, no auth\nENTSO-E API \u2014 free token",
        "\u2014",
        "FREE",
    ],
    [
        "13",
        "mFRR Energy Activation Price",
        "ENTSO-E API (A84 doc type) \u2014 free token\nRegelleistung.net \u2014 XLSX\nNetztransparenz.de \u2014 OAuth",
        "\u2014",
        "FREE (3 sources)",
    ],
]

# ---------------------------------------------------------------------------

NL_HEADERS = [
    "#",
    "Data Field",
    "Automated API Source (Free)",
    "Automated API Source (Paid)",
    "Subscription Tiers & Cost",
]

NL_ROWS = [
    [
        "1",
        "EPEX Day-Ahead Price (NL, 15-min)",
        "ENTSO-E API \u2014 free token\nEnergy-Charts API \u2014 no auth\nSMARD API (filter 252) \u2014 no auth",
        "\u2014",
        "FREE (3 independent sources)",
    ],
    [
        "2",
        "EPEX Intraday Auction 1 (IDA1) Price",
        "None",
        "EPEX SPOT sFTP\nModo Energy API",
        "EPEX SPOT: quote required\nModo Data Pro: ~\u20ac8,000\u201315,000/yr",
    ],
    [
        "3",
        "EPEX Intraday Auction 2 (IDA2) Price",
        "None",
        "EPEX SPOT sFTP\nModo Energy API",
        "EPEX SPOT: quote required\nModo Data Pro: ~\u20ac8,000\u201315,000/yr",
    ],
    [
        "4",
        "EPEX Intraday Auction 3 (IDA3) Price",
        "None",
        "EPEX SPOT sFTP\nModo Energy API",
        "EPEX SPOT: quote required\nModo Data Pro: ~\u20ac8,000\u201315,000/yr",
    ],
    [
        "5",
        "Intraday Continuous (IDC) Price",
        "None (possible delayed aggregates via SMARD)",
        "EPEX SPOT sFTP (real-time)\nModo Energy API",
        "Delayed: possibly FREE via SMARD\nReal-time: EPEX quote required / Modo ~\u20ac8,000\u201315,000/yr",
    ],
    [
        "6",
        "Imbalance / Settlement Price",
        "TenneT API \u2014 free, registration + token\nENTSO-E API (A85 doc type) \u2014 free token",
        "\u2014",
        "FREE (TenneT is authoritative source)",
    ],
    [
        "7",
        "FCR Capacity Price",
        "Regelleistung.net \u2014 HTTP/XLSX, no auth (same pan-EU auction as DE)\nENTSO-E API \u2014 free token",
        "\u2014",
        "FREE (note: same FCR Cooperation auction result as Germany)",
    ],
    [
        "8",
        "aFRR Positive Capacity Price",
        "TenneT API (Merit Order List) \u2014 free, token\nENTSO-E API \u2014 free token",
        "\u2014",
        "FREE",
    ],
    [
        "9",
        "aFRR Negative Capacity Price",
        "TenneT API (Merit Order List) \u2014 free, token\nENTSO-E API \u2014 free token",
        "\u2014",
        "FREE",
    ],
    [
        "10",
        "aFRR Energy Activation Price",
        "TenneT API (FRR Activations + Settlement) \u2014 free, token\nENTSO-E API (A84) \u2014 free token",
        "\u2014",
        "FREE (via PICASSO platform since Oct 2024)",
    ],
    [
        "11",
        "mFRR Capacity Price",
        "TenneT API (Merit Order List) \u2014 free, token\nENTSO-E API (A47) \u2014 free token",
        "\u2014",
        "FREE (via MARI platform since Dec 2025)",
    ],
    [
        "12",
        "mFRR Energy Activation Price",
        "TenneT API (FRR Activations) \u2014 free, token\nENTSO-E API (A84) \u2014 free token",
        "\u2014",
        "FREE (via MARI platform since Dec 2025)",
    ],
]


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # Landscape orientation for wide tables
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        # Swap width/height for landscape
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # === Title ===
    title = doc.add_heading("BESS Dashboard \u2014 Market Rate Data Subscriptions", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    meta = doc.add_paragraph()
    for lbl, val in [
        ("Purpose", "Subscription requirements for automated (API) market data ingestion"),
        ("Markets", "United Kingdom (GB), Germany (DE), Netherlands (NL)"),
        ("Prepared", "February 2026"),
    ]:
        r = meta.add_run(f"{lbl}: ")
        r.bold = True
        r.font.size = Pt(10)
        r2 = meta.add_run(f"{val}\n")
        r2.font.size = Pt(10)

    # =====================================================================
    # TABLE 1 \u2014 UK
    # =====================================================================
    add_heading(doc, "Table 1: United Kingdom (GB)", level=1)

    p = doc.add_paragraph()
    r = p.add_run("Settlement period: ")
    r.bold = True
    r.font.size = Pt(9)
    r2 = p.add_run("30 minutes (48 periods/day)  |  ")
    r2.font.size = Pt(9)
    r3 = p.add_run("Currency: ")
    r3.bold = True
    r3.font.size = Pt(9)
    r4 = p.add_run("GBP  |  ")
    r4.font.size = Pt(9)
    r5 = p.add_run("TSO: ")
    r5.bold = True
    r5.font.size = Pt(9)
    r6 = p.add_run("NESO  |  ")
    r6.font.size = Pt(9)
    r7 = p.add_run("Exchange: ")
    r7.bold = True
    r7.font.size = Pt(9)
    r8 = p.add_run("EPEX SPOT / Nord Pool")
    r8.font.size = Pt(9)

    add_table(doc, UK_HEADERS, UK_ROWS)

    # UK notes
    doc.add_paragraph()
    notes_heading = doc.add_paragraph()
    rn = notes_heading.add_run("Notes:")
    rn.bold = True
    rn.font.size = Pt(9)

    uk_notes = [
        "Rows 1\u20134 (EPEX wholesale prices) have NO free automated API. The GridBeyond monthly report is the current source at no extra cost (covered by the 5% aggregator revenue share). For live/daily automated access, a paid subscription to EPEX SPOT or Modo Energy is required.",
        "Row 5 (DA HH Reference / Market Index) is available free via Elexon BMRS, but it is a volume-weighted composite index from multiple short-term products \u2014 it is NOT identical to the raw EPEX Day-Ahead auction clearing price. It is a reasonable proxy for analytical purposes.",
        "Modo Energy \u2018Data Pro\u2019 tier (~\u00a38,000\u201315,000/yr) provides automated API access to all EPEX wholesale prices (rows 1\u20134) via a single subscription. Higher tiers (Business ~\u00a315,000\u201330,000/yr, Enterprise ~\u00a330,000\u201360,000/yr) add benchmarking and asset-level analytics.",
        "NESO Data Portal rate limit is 2 requests/minute for datastore queries \u2014 significantly slower than Elexon BMRS (5,000 req/min). Plan batch schedules accordingly.",
    ]
    for n in uk_notes:
        p = doc.add_paragraph(style="List Bullet")
        p.text = n
        for run in p.runs:
            run.font.size = Pt(8)

    # =====================================================================
    # TABLE 2 \u2014 GERMANY
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, "Table 2: Germany (DE)", level=1)

    p = doc.add_paragraph()
    r = p.add_run("Settlement period: ")
    r.bold = True
    r.font.size = Pt(9)
    r2 = p.add_run("15 minutes (96 periods/day)  |  ")
    r2.font.size = Pt(9)
    r3 = p.add_run("Currency: ")
    r3.bold = True
    r3.font.size = Pt(9)
    r4 = p.add_run("EUR  |  ")
    r4.font.size = Pt(9)
    r5 = p.add_run("TSOs: ")
    r5.bold = True
    r5.font.size = Pt(9)
    r6 = p.add_run("50Hertz, Amprion, TenneT, TransnetBW  |  ")
    r6.font.size = Pt(9)
    r7 = p.add_run("Exchange: ")
    r7.bold = True
    r7.font.size = Pt(9)
    r8 = p.add_run("EPEX SPOT")
    r8.font.size = Pt(9)

    add_table(doc, DE_HEADERS, DE_ROWS)

    doc.add_paragraph()
    notes_heading = doc.add_paragraph()
    rn = notes_heading.add_run("Notes:")
    rn.bold = True
    rn.font.size = Pt(9)

    de_notes = [
        "Day-ahead prices (row 1) are available from 3 independent free sources: SMARD (easiest \u2014 no auth), ENTSO-E (needs free token), and Energy-Charts (no auth). All source from the same EPEX auction results.",
        "IDA1/IDA2/IDA3 prices (rows 2\u20134) and real-time IDC (row 5) require a paid EPEX SPOT or Modo Energy subscription. There is no free automated source for clean intraday auction results.",
        "Regelleistung.net endpoints for FCR/aFRR/mFRR data (rows 7\u201313) are free and functional but UNDOCUMENTED \u2014 they return XLSX files via HTTP GET. These endpoints could change without notice. Build with fallback to ENTSO-E.",
        "Germany uses pay-as-bid for aFRR/mFRR (unlike UK\u2019s pay-as-cleared). Both capacity AND energy activation payments apply \u2014 the energy activation price (rows 10, 13) is a concept that does not exist in the UK model.",
        "Python libraries: entsoe-py (ENTSO-E), de-smard or deutschland[smard] (SMARD), custom requests code (Regelleistung, Netztransparenz).",
    ]
    for n in de_notes:
        p = doc.add_paragraph(style="List Bullet")
        p.text = n
        for run in p.runs:
            run.font.size = Pt(8)

    # =====================================================================
    # TABLE 3 \u2014 NETHERLANDS
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, "Table 3: Netherlands (NL)", level=1)

    p = doc.add_paragraph()
    r = p.add_run("Settlement period: ")
    r.bold = True
    r.font.size = Pt(9)
    r2 = p.add_run("15 minutes (96 periods/day)  |  ")
    r2.font.size = Pt(9)
    r3 = p.add_run("Currency: ")
    r3.bold = True
    r3.font.size = Pt(9)
    r4 = p.add_run("EUR  |  ")
    r4.font.size = Pt(9)
    r5 = p.add_run("TSO: ")
    r5.bold = True
    r5.font.size = Pt(9)
    r6 = p.add_run("TenneT NL  |  ")
    r6.font.size = Pt(9)
    r7 = p.add_run("Exchange: ")
    r7.bold = True
    r7.font.size = Pt(9)
    r8 = p.add_run("EPEX SPOT (formerly APX)")
    r8.font.size = Pt(9)

    add_table(doc, NL_HEADERS, NL_ROWS)

    doc.add_paragraph()
    notes_heading = doc.add_paragraph()
    rn = notes_heading.add_run("Notes:")
    rn.bold = True
    rn.font.size = Pt(9)

    nl_notes = [
        "Day-ahead prices (row 1) are available from 3 free sources: ENTSO-E (primary, needs free token), Energy-Charts (no auth), and SMARD (filter 252, no auth).",
        "IDA1/IDA2/IDA3 and real-time IDC (rows 2\u20135) require paid EPEX SPOT or Modo Energy subscription \u2014 same as UK and DE.",
        "TenneT API (developer.tennet.eu) is the authoritative source for NL settlement prices, balance delta, and ancillary service data. Free with registration. Python library: tenneteu-py.",
        "FCR clearing price (row 7) comes from the same pan-European FCR Cooperation auction as Germany. The Regelleistung.net endpoint returns the same results. 30% of TenneT\u2019s FCR must be sourced from NL-connected providers.",
        "The Dutch market explicitly rewards \u2018passive balancing\u2019 \u2014 TenneT publishes near-real-time balance delta signals that can be used to profit from imbalance pricing. This is a distinct NL revenue strategy not present in UK or DE.",
        "Python libraries: tenneteu-py (TenneT), entsoe-py (ENTSO-E). Same Regelleistung custom code as DE for FCR.",
    ]
    for n in nl_notes:
        p = doc.add_paragraph(style="List Bullet")
        p.text = n
        for run in p.runs:
            run.font.size = Pt(8)

    # =====================================================================
    # APPENDIX: PAID SUBSCRIPTION SUMMARY
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, "Appendix: Paid Subscription Options (All Markets)", level=1)

    p = doc.add_paragraph(
        "The table below summarises all paid subscriptions that would be needed "
        "to fill the gaps where no free automated API exists (primarily EPEX intraday auction and continuous data)."
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].italic = True

    add_table(
        doc,
        ["Provider", "Product", "Markets Covered", "Data Included", "Annual Cost (approx.)"],
        [
            [
                "EPEX SPOT",
                "GB Market Data Package",
                "UK",
                "DA 60-min, DA 30-min, IDA1, IDC (sFTP or API)",
                "Quote required (contact marketdata.sales@epexspot.com)",
            ],
            [
                "EPEX SPOT",
                "DE Market Data Package",
                "Germany",
                "DA 15-min, IDA1/2/3, IDC (sFTP or API)",
                "Quote required",
            ],
            [
                "EPEX SPOT",
                "NL Market Data Package",
                "Netherlands",
                "DA 15-min, IDA1/2/3, IDC (sFTP or API)",
                "Quote required",
            ],
            [
                "Nord Pool",
                "N2EX Day-Ahead",
                "UK",
                "N2EX DA prices + 1 API account",
                "~\u20ac1,200/yr",
            ],
            [
                "Nord Pool",
                "DE Day-Ahead",
                "Germany",
                "DE DA prices + 1 API account",
                "~\u20ac1,200/yr",
            ],
            [
                "Nord Pool",
                "FR/NL/AT/BE DA Bundle",
                "Netherlands (+ FR, AT, BE)",
                "NL DA prices + 1 API account",
                "~\u20ac1,700\u20133,300/yr",
            ],
            [
                "Nord Pool",
                "Intraday Trades (per market)",
                "UK, DE, or NL",
                "Intraday continuous trade data",
                "~\u20ac5,000/yr per market",
            ],
            [
                "Nord Pool",
                "Additional API Account",
                "Any",
                "Extra API login",
                "~\u20ac350/yr per account",
            ],
            [
                "Modo Energy",
                "Data Pro",
                "GB (primary), DE/NL (expanding)",
                "All EPEX wholesale prices via API, BESS indices",
                "~\u00a38,000\u201315,000/yr",
            ],
            [
                "Modo Energy",
                "Business",
                "GB, DE, NL",
                "Data Pro + asset-level benchmarking",
                "~\u00a315,000\u201330,000/yr",
            ],
            [
                "Modo Energy",
                "Enterprise",
                "GB, DE, NL",
                "Full platform access + custom analytics",
                "~\u00a330,000\u201360,000/yr",
            ],
        ],
    )

    doc.add_paragraph()

    # Cost scenarios
    add_heading(doc, "Cost Scenarios", level=2)

    add_table(
        doc,
        ["Scenario", "UK", "Germany", "Netherlands", "Total (approx.)"],
        [
            [
                "Free sources only (DA + ancillary + imbalance)",
                "\u00a30",
                "\u20ac0",
                "\u20ac0",
                "\u00a30 / \u20ac0",
            ],
            [
                "Free + GridBeyond report (UK only)",
                "\u00a30 (5% revenue share)",
                "N/A",
                "N/A",
                "\u00a30",
            ],
            [
                "Add EPEX intraday data (all 3 markets)",
                "EPEX quote",
                "EPEX quote",
                "EPEX quote",
                "Quote required",
            ],
            [
                "Modo Energy Data Pro (covers EPEX gap)",
                "~\u00a38,000\u201315,000",
                "Included if covered",
                "Included if covered",
                "~\u00a38,000\u201315,000/yr",
            ],
            [
                "Nord Pool (DA + Intraday, all markets)",
                "~\u20ac6,200",
                "~\u20ac6,200",
                "~\u20ac6,700\u20138,300",
                "~\u20ac19,100\u201320,700/yr (~\u00a316,000\u201317,500)",
            ],
            [
                "Full (Modo Business + Nord Pool backup)",
                "~\u00a321,000\u201336,000",
                "~\u20ac6,200",
                "~\u20ac6,700\u20138,300",
                "~\u00a333,000\u201350,000/yr",
            ],
        ],
    )

    # Save
    output_path = r"C:\repos\bess-dashboard\docs\Market_Data_Subscriptions_Tables.docx"
    doc.save(output_path)
    print(f"Document saved: {output_path}")


if __name__ == "__main__":
    main()
